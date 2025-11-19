# core/docx_writer.py
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import os


def add_table_borders(table):
    """í‘œì— í…Œë‘ë¦¬ ì¶”ê°€"""
    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)


def set_cell_background(cell, color):
    """ì…€ ë°°ê²½ìƒ‰ ì„¤ì •"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)


def parse_user_query(user_query: str) -> dict:
    """
    ì‚¬ìš©ì ì¿¼ë¦¬ì—ì„œ ê³µì‚¬ëª… / ì‚¬ê³ ë°œìƒì¥ì†Œ / ì‚¬ê³ ì¢…ë¥˜ / ì‚¬ê³ ê°œìš” ì¶”ì¶œ
    """
    data = {
        "ê³µì‚¬ëª…": "",
        "ì‚¬ê³ ë°œìƒì¥ì†Œ": "",
        "ì‚¬ê³ ì¢…ë¥˜": "",
        "ì‚¬ê³ ê°œìš”": ""
    }
    
    if not user_query:
        return data
    
    lines = user_query.strip().split('\n')
    for line in lines:
        line = line.strip()
        if 'ê³µì¢…:' in line:
            data["ê³µì‚¬ëª…"] = line.split(':', 1)[1].strip()
        elif 'ì‘ì—…í”„ë¡œì„¸ìŠ¤' in line:
            data["ì‚¬ê³ ë°œìƒì¥ì†Œ"] = line.split(':', 1)[1].strip()
        elif 'ì‚¬ê³  ìœ í˜•' in line or 'ì‚¬ê³ ìœ í˜•' in line:
            data["ì‚¬ê³ ì¢…ë¥˜"] = line.split(':', 1)[1].strip()
        elif 'ì‚¬ê³  ê°œìš”' in line:
            data["ì‚¬ê³ ê°œìš”"] = line.split(':', 1)[1].strip()
    
    return data


def _fill_multiline_text(cell, text: str, font_size: int = 9):
    """
    ì…€ì— ì—¬ëŸ¬ ì¤„ í…ìŠ¤íŠ¸(ì—¬ëŸ¬ ë¬¸ë‹¨)ë¥¼ ì±„ì›Œ ë„£ê¸° ìœ„í•œ ìœ í‹¸.
    ê¸°ì¡´ ë‚´ìš© ì´ˆê¸°í™” í›„ ì¤„ë§ˆë‹¤ ìƒˆ paragraph ìƒì„±.
    """
    cell.text = ""
    if not text:
        return
    
    lines = text.split('\n')
    for line in lines:
        para = cell.add_paragraph(line.strip())
        para.paragraph_format.line_spacing = 1.3
        para.paragraph_format.space_after = Pt(4)
        for run in para.runs:
            run.font.size = Pt(font_size)


def create_accident_report_docx(
    user_query: str,
    cause_text: str,
    action_text: str,
    output_path: str = None,
    source_references: list = None,  # âœ… ì¶”ê°€
) -> str:
    """
    [ë³„ì§€ ì œ2í˜¸ ì„œì‹] ê±´ì„¤ì‚¬ê³  ë°œìƒí˜„í™© ë³´ê³  ì–‘ì‹ DOCX ìƒì„±
    + ê´€ë ¨ ê·¼ê±° ìë£Œ ì¶”ê°€ (ì„ íƒì )
    """
    doc = Document()
    
    # ==== í˜ì´ì§€ ì—¬ë°± ====
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # ==== í—¤ë” ====
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header_para.add_run('[ë³„ì§€ ì œ2í˜¸ ì„œì‹] ê±´ì„¤ì‚¬ê³  ë°œìƒí˜„í™© ë³´ê³  ì–‘ì‹')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(100, 100, 100)
    
    # ==== ì œëª© ====
    title = doc.add_heading('ê±´ì„¤ì‚¬ê³  ë°œìƒí˜„í™© ë³´ê³ ', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.name = 'ë§‘ì€ ê³ ë”•'
    
    doc.add_paragraph()
    
    # ==== ì‚¬ìš©ì ì§ˆì˜ íŒŒì‹± ====
    query_data = parse_user_query(user_query)
    
    # ==== ê¸°ë³¸ì •ë³´ í…Œì´ë¸” ====
    table1 = doc.add_table(rows=3, cols=4)
    table1.style = 'Table Grid'
    add_table_borders(table1)
    
    # 1í–‰
    cells = table1.rows[0].cells
    set_cell_background(cells[0], 'E7E6E6')
    set_cell_background(cells[2], 'E7E6E6')
    cells[0].text = 'ìˆ˜ì‹ '
    cells[1].text = ''
    cells[2].text = 'ë³´ê³ ì¼ì‹œ'
    cells[3].text = datetime.now().strftime('%Yë…„ %mì›” %dì¼ %Hì‹œ %Më¶„')
    
    for cell in [cells[0], cells[2]]:
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.bold = True
                run.font.size = Pt(10)
    
    # 2í–‰
    cells = table1.rows[1].cells
    set_cell_background(cells[0], 'E7E6E6')
    set_cell_background(cells[2], 'E7E6E6')
    cells[0].text = 'ë°œì‹ ê¸°ê´€'
    cells[1].text = ''
    cells[2].text = 'ë°œì‹ (ë³´ê³ ì)'
    cells[3].text = 'O O O (ì¸)'
    
    # 3í–‰
    cells = table1.rows[2].cells
    set_cell_background(cells[0], 'E7E6E6')
    cells[0].text = 'ì œëª©'
    cells[1].merge(cells[2]).merge(cells[3])
    cells[1].text = 'ê±´ì„¤ì‚¬ê³  ë°œìƒí˜„í™© ë³´ê³ '
    
    doc.add_paragraph()
    
    # ==== ì‚¬ê³  ìƒì„¸ ì •ë³´ í…Œì´ë¸” ====
    table2 = doc.add_table(rows=15, cols=4)  # âœ… 15í–‰ìœ¼ë¡œ ë³€ê²½ (14í–‰ê¹Œì§€ ì‚¬ìš© + ë¹„ê³  1í–‰)
    table2.style = 'Table Grid'
    add_table_borders(table2)
    
    row_data = [
        (0, 'ì‚¬ê³ ì¼ì‹œ', datetime.now().strftime('%Yë…„ %mì›” %dì¼ ( )ìš”ì¼  ì‹œ  ë¶„ ê²½'), [(2, 'ê¸°ìƒìƒíƒœ', '')]),
        (1, 'ê³µì‚¬ëª…', query_data.get('ê³µì‚¬ëª…', ''), None),
        (2, 'ì‹œê³µì‚¬', 'ì±…ì„ì ë° ì—°ë½ì²˜', None),
        (3, 'ê±´ì„¤ì‚¬ì—…ê´€ë¦¬ê¸°ìˆ ì', 'ì±…ì„ì ë° ì—°ë½ì²˜', None),
        (4, 'ì„¤ê³„ì', 'ì±…ì„ì ë° ì—°ë½ì²˜', None),
        (5, 'í˜„ì¥ì£¼ì†Œ', '', [(2, 'ì‚¬ê³ ë°œìƒì¥ì†Œ', query_data.get('ì‚¬ê³ ë°œìƒì¥ì†Œ', ''))]),
        (6, 'ì‚¬ê³  ì¢…ë¥˜', query_data.get('ì‚¬ê³ ì¢…ë¥˜', ''), None),
        (7, 'ì¸ì í”¼í•´', '', [(2, 'ì¥ë¹„ì†ì‹¤', '')]),
        (8, 'êµ¬ì¡°ë¬¼ ì†ì‹¤', '', [(2, 'í”¼í•´ê¸ˆì•¡', '')]),
        (9, 'ê³µê¸°ì§€ì—°', '', [(2, 'ì•ˆì „ê´€ë¦¬ê³„íšì„œ\nìˆ˜ë¦½ì—¬ë¶€', 'í•´ë‹¹ : (  ), í•´ë‹¹ì—†ìŒ : (  )\ní•´ë‹¹ì‚¬ìœ  : ì˜ ì œ98ì¡°ì œ1í•­(  )í˜¸')]),
    ]
    
    for row_idx, label, value, extra in row_data:
        cells = table2.rows[row_idx].cells
        
        set_cell_background(cells[0], 'E7E6E6')
        cells[0].text = label
        cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cells[0].paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(9)
        
        if extra:
            cells[1].text = value
            set_cell_background(cells[2], 'E7E6E6')
            cells[2].text = extra[0][1]
            for run in cells[2].paragraphs[0].runs:
                run.font.bold = True
                run.font.size = Pt(9)
            cells[3].text = extra[0][2]
        else:
            cells[1].merge(cells[2]).merge(cells[3])
            cells[1].text = value
    
    # ==== ì‚¬ê³ ë°œìƒ ê²½ìœ„ ====
    row_idx = 10
    cells = table2.rows[row_idx].cells
    set_cell_background(cells[0], 'E7E6E6')
    cells[0].text = 'ì‚¬ê³ ë°œìƒ ê²½ìœ„\n(ë°œìƒì›ì¸)'
    _fill_multiline_text(cells[1].merge(cells[2]).merge(cells[3]), cause_text)
    
    # ==== ì¡°ì¹˜ì‚¬í•­ ë° í–¥í›„ì¡°ì¹˜ê³„íš ====
    row_idx = 11
    cells = table2.rows[row_idx].cells
    set_cell_background(cells[0], 'E7E6E6')
    cells[0].text = 'ì¡°ì¹˜ì‚¬í•­ ë°\ní–¥í›„ì¡°ì¹˜ê³„íš'
    _fill_multiline_text(cells[1].merge(cells[2]).merge(cells[3]), action_text)
    
    # ==== ì‚¬ê³ ì¡°ì‚¬ ë°©ë²• ====
    row_idx = 12
    cells = table2.rows[row_idx].cells
    set_cell_background(cells[0], 'E7E6E6')
    cells[0].text = 'ì‚¬ê³ ì¡°ì‚¬ ë°©ë²•'
    cells[1].merge(cells[2]).merge(cells[3]).text = "1. ì§ì ‘ì¡°ì‚¬\n2. ì‚¬ê³ ì¡°ì‚¬ìœ„ì›íšŒì¡°ì‚¬\n3. ë…¸ë™ë¶€ ì¬í•´ì¡°ì‚¬ì‹œ í•©ë™ì¡°ì‚¬"
    
    # âœ… ==== ê´€ë ¨ ê·¼ê±° ìë£Œ (13í–‰) ====
    row_idx = 13
    if source_references and len(source_references) > 0:
        cells = table2.rows[row_idx].cells
        set_cell_background(cells[0], 'E7E6E6')
        cells[0].text = 'ê´€ë ¨\nê·¼ê±°ìë£Œ'
        cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cells[0].paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(9)
        
        # ê·¼ê±° ìë£Œ ë‚´ìš© ì‘ì„±
        merged_cell = cells[1].merge(cells[2]).merge(cells[3])
        merged_cell.text = ""  # ì´ˆê¸°í™”
        
        # ê° ê·¼ê±° ìë£Œë¥¼ ë¬¸ë‹¨ìœ¼ë¡œ ì¶”ê°€
        for i, ref in enumerate(source_references, 1):
            # êµ¬ë¶„ì„  (ì²« ë²ˆì§¸ ì œì™¸)
            if i > 1:
                sep_para = merged_cell.add_paragraph()
                sep_run = sep_para.add_run('â”€' * 60)
                sep_run.font.size = Pt(8)
                sep_run.font.color.rgb = RGBColor(180, 180, 180)
            
            # ë¬¸ì„œ ì œëª©
            title_para = merged_cell.add_paragraph()
            title_run = title_para.add_run(f'[ë¬¸ì„œ {ref["idx"]}] {ref["filename"]}')
            title_run.font.size = Pt(9)
            title_run.font.bold = True
            title_run.font.color.rgb = RGBColor(0, 70, 140)
            title_para.paragraph_format.space_after = Pt(2)
            
            # ìœ„ì¹˜ ì •ë³´
            if ref.get("hierarchy"):
                loc_para = merged_cell.add_paragraph()
                loc_run1 = loc_para.add_run('ğŸ“ ìœ„ì¹˜: ')
                loc_run1.font.size = Pt(8)
                loc_run1.font.bold = True
                loc_run2 = loc_para.add_run(ref["hierarchy"])
                loc_run2.font.size = Pt(8)
                loc_para.paragraph_format.space_after = Pt(2)
            elif ref.get("section"):
                loc_para = merged_cell.add_paragraph()
                loc_run1 = loc_para.add_run('ğŸ“ ì„¹ì…˜: ')
                loc_run1.font.size = Pt(8)
                loc_run1.font.bold = True
                loc_run2 = loc_para.add_run(ref["section"])
                loc_run2.font.size = Pt(8)
                loc_para.paragraph_format.space_after = Pt(2)
            
            # ê´€ë ¨ì„± ìš”ì•½
            if ref.get("relevance_summary"):
                rel_para = merged_cell.add_paragraph()
                rel_run1 = rel_para.add_run('ğŸ’¡ ê´€ë ¨ì„±: ')
                rel_run1.font.size = Pt(8)
                rel_run1.font.bold = True
                rel_run2 = rel_para.add_run(ref["relevance_summary"])
                rel_run2.font.size = Pt(8)
                rel_para.paragraph_format.space_after = Pt(2)
            
            # í•µì‹¬ ë‚´ìš©
            key_sentences = ref.get("key_sentences", [])
            if key_sentences:
                key_para = merged_cell.add_paragraph()
                key_run = key_para.add_run('ğŸ¯ í•µì‹¬ ë‚´ìš©:')
                key_run.font.size = Pt(8)
                key_run.font.bold = True
                key_para.paragraph_format.space_after = Pt(2)
                
                for sentence in key_sentences[:3]:  # ìµœëŒ€ 3ê°œë§Œ
                    sent_para = merged_cell.add_paragraph()
                    sent_para.paragraph_format.left_indent = Cm(0.5)
                    sent_run = sent_para.add_run(f'â€¢ {sentence}')
                    sent_run.font.size = Pt(8)
                    sent_para.paragraph_format.space_after = Pt(1)
    else:
        # ê·¼ê±° ìë£Œê°€ ì—†ìœ¼ë©´ ë¹„ê³ ë¡œ ì‚¬ìš©
        cells = table2.rows[row_idx].cells
        set_cell_background(cells[0], 'E7E6E6')
        cells[0].text = 'ë¹„ê³ '
        cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cells[0].paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(9)
        cells[1].merge(cells[2]).merge(cells[3]).text = ""
    
    # âœ… ==== ë¹„ê³  (14í–‰, ë§¨ ë§ˆì§€ë§‰) ====
    row_idx = 14
    cells = table2.rows[row_idx].cells
    set_cell_background(cells[0], 'E7E6E6')
    cells[0].text = 'ë¹„ê³ '
    cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cells[0].paragraphs[0].runs:
        run.font.bold = True
        run.font.size = Pt(9)
    cells[1].merge(cells[2]).merge(cells[3]).text = ""
    
    # ==== íŒŒì¼ ì €ì¥ ====
    if not os.path.exists("reports"):
        os.makedirs("reports")
    
    if output_path is None:
        output_path = f"reports/ê±´ì„¤ì‚¬ê³ _ë°œìƒí˜„í™©_ë³´ê³ _{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    
    doc.save(output_path)
    return output_path